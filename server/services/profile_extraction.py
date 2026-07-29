"""Extract profile proposals from private source documents with explicit consent."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from server.adapters.llm_provider import LLMProvider
from server.models.entities import (
    CandidateProfile,
    FileObject,
    ModelCredential,
    ProfileProposal,
    SourceDocument,
    User,
    utc_now,
)


class ProfileExtractionError(ValueError):
    pass


class ConsentRequired(ProfileExtractionError):
    pass


class SourceDocumentNotFound(ProfileExtractionError):
    pass


class ProfileService:
    PROFILE_SCHEMA = {
        "name": "string",
        "contact": "object",
        "education": "array",
        "experience": "array",
        "projects": "array",
        "skills": "array",
        "target": "object",
        "preferences": "object",
    }

    def __init__(self, db: Session, storage, llm_provider: LLMProvider):
        self.db = db
        self.storage = storage
        self.llm_provider = llm_provider

    def set_ai_consent(self, user_id: str, enabled: bool) -> User:
        user = self.db.get(User, user_id)
        if user is None:
            raise ProfileExtractionError("user not found")
        user.ai_processing_enabled = enabled
        user.ai_consent_at = utc_now() if enabled else None
        self.db.commit()
        return user

    def get_ai_consent(self, user_id: str) -> User:
        user = self.db.get(User, user_id)
        if user is None:
            raise ProfileExtractionError("user not found")
        return user

    def create_proposal(self, user_id: str, document_ids: list[str]) -> ProfileProposal:
        user = self.db.get(User, user_id)
        if user is None:
            raise ProfileExtractionError("user not found")
        if not user.ai_processing_enabled:
            raise ConsentRequired("AI processing consent is required")
        if not document_ids:
            raise ProfileExtractionError("at least one source document is required")

        rows = list(
            self.db.execute(
                select(SourceDocument, FileObject)
                .join(FileObject, SourceDocument.file_id == FileObject.id)
                .where(
                    SourceDocument.user_id == user_id,
                    SourceDocument.id.in_(document_ids),
                )
            ).all()
        )
        if len(rows) != len(set(document_ids)):
            raise SourceDocumentNotFound("one or more source documents were not found")

        source_parts: list[str] = []
        source_refs: list[str] = []
        for _document, file_object in rows:
            source_parts.append(self._extract_text(file_object))
            source_refs.append(file_object.filename)

        provider = self._provider_for_user(user_id)
        proposed = provider.extract_profile("\n\n".join(source_parts), self.PROFILE_SCHEMA)
        if not isinstance(proposed, dict):
            raise ProfileExtractionError("model returned an invalid profile proposal")
        model_refs = proposed.get("source_refs", [])
        if not isinstance(model_refs, list):
            model_refs = []
        merged_refs = list(dict.fromkeys([*source_refs, *(str(ref) for ref in model_refs)]))

        proposal = ProfileProposal(
            user_id=user_id,
            status="pending",
            proposed_data=proposed,
            source_refs=merged_refs,
        )
        self.db.add(proposal)
        self.db.commit()
        return proposal

    def confirm_proposal(
        self,
        user_id: str,
        proposal_id: str,
        accepted_fields: list[str],
    ) -> CandidateProfile:
        proposal = self.db.scalar(
            select(ProfileProposal).where(
                ProfileProposal.id == proposal_id,
                ProfileProposal.user_id == user_id,
            )
        )
        if proposal is None:
            raise ProfileExtractionError("profile proposal not found")
        if proposal.status != "pending":
            raise ProfileExtractionError("profile proposal is no longer pending")

        accepted = list(dict.fromkeys(accepted_fields))
        unknown = [field for field in accepted if field not in proposal.proposed_data]
        if unknown:
            raise ProfileExtractionError("accepted field is not present in the proposal")
        data = {field: proposal.proposed_data[field] for field in accepted}
        latest_version = self.db.scalar(
            select(func.max(CandidateProfile.version)).where(CandidateProfile.user_id == user_id)
        ) or 0
        profile = CandidateProfile(
            user_id=user_id,
            version=int(latest_version) + 1,
            status="confirmed",
            data=data,
            source_refs=list(proposal.source_refs),
            confirmed_at=utc_now(),
        )
        proposal.status = "confirmed"
        proposal.accepted_fields = accepted
        proposal.confirmed_at = profile.confirmed_at
        self.db.add(profile)
        self.db.commit()
        return profile

    def get_confirmed_profile(self, user_id: str) -> CandidateProfile | None:
        return self.db.scalars(
            select(CandidateProfile)
            .where(
                CandidateProfile.user_id == user_id,
                CandidateProfile.status == "confirmed",
            )
            .order_by(CandidateProfile.version.desc())
        ).first()

    def _provider_for_user(self, user_id: str) -> LLMProvider:
        provider = self.llm_provider
        if not hasattr(provider, "for_user"):
            return provider
        credential = self.db.scalar(
            select(ModelCredential)
            .where(
                ModelCredential.user_id == user_id,
                ModelCredential.enabled.is_(True),
            )
        )
        return provider.for_user(credential.encrypted_key if credential else None)

    def _extract_text(self, file_object: FileObject) -> str:
        content = self.storage.open(file_object.object_key).read()
        suffix = Path(file_object.filename).suffix.lower()
        if suffix in {".txt", ".md", ".markdown"}:
            return content.decode("utf-8", errors="replace")
        if suffix == ".pdf":
            return "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
        if suffix == ".docx":
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            paragraphs.extend(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )
            return "\n".join(paragraphs)
        raise ProfileExtractionError("unsupported source document format")
