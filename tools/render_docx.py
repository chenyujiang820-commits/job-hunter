"""DOCX 生成 — 将 Markdown 简历和求职信渲染为 .docx 文件。

支持中文排版，基础格式：
- 一级标题 → 16pt 加粗
- 二级标题 → 14pt 加粗
- 正文 → 11pt
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _set_font(run, name: str = "微软雅黑", size: int = 11, bold: bool = False) -> None:
    """设置 run 的字体属性。"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    # 设置中文字体
    run._element.rPr.rFonts.set("{}eastAsia", name)  # type: ignore[union-attr]


def _add_paragraph(doc, text: str, size: int = 11, bold: bool = False) -> None:
    """添加段落。"""
    if not text.strip():
        return
    # 清理 XML 非法字符
    import re
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text.strip())
    if not cleaned:
        return
    para = doc.add_paragraph()
    run = para.add_run(cleaned)
    _set_font(run, size=size, bold=bold)


def _add_section_title(doc, text: str) -> None:
    """添加章节标题（加粗、大字号）。"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_font(run, size=14, bold=True)
    para.space_before = Pt(12)
    para.space_after = Pt(4)


def _render_markdown_to_docx(doc: Document, markdown: str) -> None:
    """将 Markdown 文本渲染到 Document 对象。

    支持的 Markdown 元素:
    - # 一级标题 → 姓名/大标题
    - ## 二级标题 → 章节标题
    - - 列表项 → 缩进段落
    - **粗体** → 加粗
    - 普通文本 → 正文
    """
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 标题
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            _set_font(run, size=16, bold=True)
            i += 1
            continue

        if line.startswith("## "):
            _add_section_title(doc, line[3:].strip())
            i += 1
            continue

        # 列表项
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.style = doc.styles["List Bullet"]
            run = para.add_run(text)
            _set_font(run, size=11)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落 — 收集连续非空行
        paragraph_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "-", "*")):
            paragraph_lines.append(lines[i].strip())
            i += 1

        if paragraph_lines:
            text = " ".join(paragraph_lines)
            _add_paragraph(doc, text, size=11)
        else:
            i += 1


def generate_docx(
    output_path: str | Path,
    resume_md: str,
    cover_letter_md: str | None = None,
    title: str = "求职材料",
) -> Path:
    """生成 DOCX 文件，包含简历和可选的求职信。

    Args:
        output_path: 输出路径（.docx）
        resume_md: Markdown 格式简历
        cover_letter_md: Markdown 格式求职信（可选）
        title: 文件标题

    Returns:
        Path: 生成的文件路径
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    _render_markdown_to_docx(doc, resume_md)

    if cover_letter_md:
        doc.add_page_break()
        _render_markdown_to_docx(doc, cover_letter_md)

    doc.save(str(output_path))
    return output_path


# 别名兼容
def render_docx(template_path, draft, output_path) -> Path:
    """host 兼容接口: render_docx(template, ApplicationDraft, path) → Path。

    将 draft.resume_sections 转为 Markdown 后渲染。
    至少需要"summary"和一个其他章节。
    """
    sections = getattr(draft, "resume_sections", {}) or {}
    if len(sections) < 2 or set(sections.keys()) == {"summary"}:
        raise ValueError("resume_sections 缺少必要章节（至少需要 summary + 一个其他章节）")

    resume_lines = []
    for title, content in sections.items():
        resume_lines.append(f"## {title}")
        resume_lines.append(str(content))
        resume_lines.append("")
    resume_md = "\n".join(resume_lines) or "# 简历"

    cover_md = getattr(draft, "cover_letter_text", "") or ""

    return generate_docx(str(output_path), resume_md, cover_md)
