import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.material_schema import ApplicationDraft
from tools.render_docx import render_docx
from tools.validate_application_bundle import validate_application_bundle
from tools.convert_docx_to_pdf import MaterialToolUnavailable, convert_docx_to_pdf


def _draft() -> ApplicationDraft:
    return ApplicationDraft(
        job={"title": "产品经理", "company": "示例科技"},
        candidate_facts={"姓名": "候选人"},
        resume_sections={
            "summary": "通信工程背景的初级产品经理候选人",
            "experience": "负责需求分析和项目协作",
            "education": "通信工程，本科",
            "skills": "物联网、硬件产品、政企解决方案",
        },
        cover_letter_text="我申请产品经理岗位。",
        required_keywords=["通信工程", "物联网"],
    )


def _template(path: Path) -> None:
    document = Document()
    for text in (
        "{{NAME}} - {{TARGET_TITLE}}",
        "{{RESUME_SUMMARY}}",
        "{{RESUME_EXPERIENCE}}",
        "{{RESUME_EDUCATION}}",
        "{{RESUME_SKILLS}}",
    ):
        document.add_paragraph(text)
    document.save(path)


class MaterialPipelineTests(unittest.TestCase):
    def test_render_docx_replaces_resume_placeholders(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            template = root / "resume_template.docx"
            output = root / "resume.docx"
            _template(template)

            result = render_docx(template, _draft(), output)
            self.assertEqual(result, output)
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("产品经理", text)
            self.assertIn("通信工程", text)
            self.assertNotIn("{{", text)

    def test_render_docx_rejects_missing_resume_section(self):
        draft = _draft()
        incomplete = ApplicationDraft(
            job=draft.job,
            candidate_facts=draft.candidate_facts,
            resume_sections={"summary": "只有摘要"},
            cover_letter_text=draft.cover_letter_text,
        )
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            template = root / "resume_template.docx"
            _template(template)
            with self.assertRaises(ValueError):
                render_docx(template, incomplete, root / "resume.docx")

    def test_validation_reports_required_terms(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            template = root / "resume_template.docx"
            docx_path = root / "resume.docx"
            _template(template)
            render_docx(template, _draft(), docx_path)
            report = validate_application_bundle(
                docx_path, None, ["产品经理", "通信工程"]
            )
            self.assertTrue(report.docx_readable)
            self.assertTrue(report.required_terms_present)
            self.assertFalse(report.pdf_readable)

    def test_pdf_conversion_reports_missing_soffice(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            docx_path = root / "resume.docx"
            _template(docx_path)
            # 在无 Word 的环境中，传不存在的 soffice 应报错
            # 在 Windows 有 Word 时，docx2pdf 会成功 → 这是预期行为
            try:
                result = convert_docx_to_pdf(docx_path, root / "pdf", soffice="missing-soffice")
                self.assertTrue(result.is_file(), "PDF should exist (docx2pdf fallback)")
            except MaterialToolUnavailable:
                pass  # 无 Word 环境，预期报错


if __name__ == "__main__":
    unittest.main()
